#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
__title__ = ''
__author__ = 'mike_jun'
__mtime__ = '2019-7-1'
#目的： 1. 先将一个文件夹下的所有文件夹的 .java 文件路径保存到一个列表中
        2. 依次读取列表的路径， 将 .java 文件内容保存到word 中
"""
import enum
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re
from docx.shared import Length


fileList = [] # 使用全局列表保存文件路径
def getAllFile(path, fileList):  # 使用递归方法
    dirList = []  # 保存文件夹
    files = os.listdir(path) # 返回一个列表，其中包含文件 和 文件夹
    for f in files:
        if (os.path.isdir(path + '/' + f)):
            dirList.append(path + '/' + f)    #  将文件夹 名字 进行保存

        if (os.path.isfile(path + '/' + f)):
            fileList.append(path + '/' + f) # 将文件名 保存

    for dir in dirList: #如果文件夹为空时，递归自动退出
        getAllFile(dir, fileList) # 递归保存到将.java 文件保存到 fileList 中

# fileList.append('/data1/private/cc/THUCC/main.py')

getAllFile('/data1/private/cc/THUCC', fileList)

new_fileList = []
for file in fileList:
    must_have = ['systems']
    must_not_have = ['transformers', 'translate', 'wsd', 'poem_appreciation']
    if not any([x in file for x in must_have]):
        continue
    if any([x in file for x in must_not_have]):
        continue
    new_fileList.append(file)
fileList = new_fileList

def getPythonFile(fileList):
    new_fileList = []
    for file in fileList:
        if file.endswith('.py') or file.endswith('.yml'):
            new_fileList.append(file)
    print('文件数量为： ',len(new_fileList))
    return new_fileList

fileList = getPythonFile(fileList)

print('文件数量为： ',len(fileList))
print(fileList)
print(os.path.isfile(fileList[0])) # 判断第一个值是否是文件

def file_analysis(old_file_lines):
    """标记需要删除的注释的行号，并存入列表"""
    hashtap = []
    six_quotes = []
    i = 0
    for line in old_file_lines:
        line = line.replace('\n', '')
        # 符号 # 独占一行
        ret_1 = re.match(r"^[^\w]*#+",line)
        if ret_1:
            hashtap.append(i)
        # 符号 """ 独占一行
        ret_2 = re.match(r"[ ]*r?\"\"\"",line)
        if ret_2:
            # 如果存在类型，函数说明的 """xxxxx""" 之类的，不予删除
            ret_2_1 = re.match(r"[^\"]*\"\"\"[^\"]*\"\"\"",line)
            if ret_2_1:
                pass
            else:
                six_quotes.append(i)
        i += 1
    # 将两个"""行号之间所有的行添加到 # 号列表中
    while six_quotes != []:
        # 从列表中移出最后两个元素
        a = six_quotes.pop()
        b = six_quotes.pop()
        temp = b
        while temp <= a:
            hashtap.append(temp)
            temp += 1
    hashtap = set(hashtap)
    new_file_lines = []
    for i, line in enumerate(old_file_lines):
        if i not in hashtap:
            if line != '\n':
                new_file_lines.append(line)
    return new_file_lines


def saveDocFile():
    # SINGLE         =>  单倍行距（默认）
    # ONE_POINT_FIVE =>  1.5倍行距
    # DOUBLE2        =>  倍行距
    # AT_LEAST       =>  最小值
    # EXACTLY        =>  固定值
    # MULTIPLE       =>  多倍行距
    doc = Document()
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph('') #增加一页
    doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
    doc.styles['Normal'].font.size = Pt(8) #  设置字体的大小为 5 号字体
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码
    save_file = r'code3.doc'
    codeNum = 0
    all_lines = []
    for i, f in enumerate(fileList):
        print('starting deal %d'%i)
        try:
            with open(f, encoding='UTF-8') as file:  # 转换编码以实现正确输出中文格式
                file_lines = file_analysis(file.readlines())
                all_lines.extend(file_lines)
        except:
            continue

    print(len(all_lines))
    if len(all_lines) > 4000:
        all_lines = all_lines[:2000] + all_lines[-2000:]
    print(len(all_lines))
    for line in all_lines:
        p.add_run(line)
    doc.save(save_file)  # 不足60 页进行保存
    print('all done')

saveDocFile()
print('all done')